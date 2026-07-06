from pathlib import Path

from docx import Document
from docx.shared import RGBColor

from csv_to_word_forms import (
    PANDADOC_SIGNATURE_ROLE,
    PANDADOC_SIGNATURE_TAG,
    build_column_map,
    collect_all_tables,
    fill_template,
    generate_combined_docx,
    get_row_values,
)

TEMPLATE_PATH = (
    Path(__file__).parent.parent
    / "Form Template"
    / "Rapid Rehousing Program Check Request Form - Template.docx"
)

SAMPLE_ROW = {
    "Client Name": "Test Client",
    "Check Type": "Monthly Rent",
    "UBH Amount": "1500",
    "Client Rent Amount": "200",
}


def _all_table_text(doc: Document) -> str:
    parts = []
    for table in collect_all_tables(doc):
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _fill(include_signature_tag: bool) -> Document:
    doc = Document(str(TEMPLATE_PATH))
    column_map = build_column_map(list(SAMPLE_ROW.keys()))
    values = get_row_values(SAMPLE_ROW, column_map)
    fill_template(
        doc, values, SAMPLE_ROW, column_map, include_signature_tag=include_signature_tag
    )
    return doc


def test_signature_tag_constants_shape():
    assert PANDADOC_SIGNATURE_ROLE
    assert PANDADOC_SIGNATURE_TAG.startswith("[")
    assert PANDADOC_SIGNATURE_TAG.endswith("]")
    assert PANDADOC_SIGNATURE_ROLE in PANDADOC_SIGNATURE_TAG


def test_fill_template_default_has_no_signature_tag():
    doc = _fill(include_signature_tag=False)
    assert PANDADOC_SIGNATURE_TAG not in _all_table_text(doc)


def test_fill_template_opt_in_writes_tag_into_director_cell():
    doc = _fill(include_signature_tag=True)

    tagged_cells = []
    for table in collect_all_tables(doc):
        for row in table.rows:
            for cell in row.cells:
                if PANDADOC_SIGNATURE_TAG in cell.text:
                    tagged_cells.append(cell)

    assert len(tagged_cells) >= 1
    assert any("Director:" in cell.text for cell in tagged_cells)


def test_signature_tag_run_is_white():
    doc = _fill(include_signature_tag=True)

    white_tag_runs = 0
    for table in collect_all_tables(doc):
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if PANDADOC_SIGNATURE_TAG in run.text:
                            assert run.font.color.rgb == RGBColor(0xFF, 0xFF, 0xFF)
                            white_tag_runs += 1
    assert white_tag_runs >= 1


def test_generate_combined_docx_from_rows(tmp_path):
    rows = [dict(SAMPLE_ROW), {**SAMPLE_ROW, "Client Name": "Second Client"}]
    out_path = tmp_path / "combined.docx"

    result = generate_combined_docx(
        rows, TEMPLATE_PATH, out_path, include_signature_tag=True
    )

    assert result == out_path
    assert out_path.is_file()
    merged = Document(str(out_path))
    text = _all_table_text(merged)
    assert text.count(PANDADOC_SIGNATURE_TAG) == 2
    assert "Test Client" in text
    assert "Second Client" in text


def test_generate_combined_docx_without_tag(tmp_path):
    out_path = tmp_path / "combined_no_tag.docx"

    generate_combined_docx([dict(SAMPLE_ROW)], TEMPLATE_PATH, out_path)

    merged = Document(str(out_path))
    assert PANDADOC_SIGNATURE_TAG not in _all_table_text(merged)


def test_generate_combined_docx_empty_rows_returns_none(tmp_path):
    out_path = tmp_path / "combined_empty.docx"

    result = generate_combined_docx([], TEMPLATE_PATH, out_path)

    assert result is None
    assert not out_path.exists()
