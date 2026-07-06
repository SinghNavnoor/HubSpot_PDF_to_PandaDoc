#!/usr/bin/env python3
"""
CSV to Word Check Request Form Generator

Reads check request data from the Data folder (the single CSV file in that folder)
and generates one combined Word document with each filled form on its own page,
plus a PDF copy (via Microsoft Word + docx2pdf, or LibreOffice headless).
Scales for any number of rows (5, 25, etc.).
"""
from __future__ import annotations

import argparse
import csv
import re
import shutil
import subprocess
import sys
import warnings
from copy import deepcopy
from datetime import date, datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_UNDERLINE
from docx.shared import Pt, RGBColor
from docxcompose.composer import Composer

# PandaDoc field tag embedded in the Director signature cell when
# include_signature_tag=True. Bracket notation per
# https://developers.pandadoc.com/docs/field-tags — [fieldType:role_____].
# The role must match the recipient role sent in the PandaDoc API request
# (Phase 3). Rendered in white so it is invisible on the printed form.
# NOTE: exact syntax to be confirmed in the PandaDoc sandbox per the design spec.
PANDADOC_SIGNATURE_ROLE = "ProgramDirector"
PANDADOC_SIGNATURE_TAG = f"[signature:{PANDADOC_SIGNATURE_ROLE}____________]"

# (CSV column header — any casing; matched via normalize_column_name, template label in Word doc)
FIELD_MAPPING = [
    ("client name", "Client Name:"),
    ("payment date (Today's Date)", "Date:"),
    ("Program (Sync)", "Program:"),
    ("check type", "Check Type:"),
    ("Check Payable to (Sync)", "Check Payable to:"),
    ("Landlord Address Sync", "Landlord Address:"),
    ("Monthly Rent Amount", "Total Monthly Rent:"),
    ("ubh amount", "UBH Rent Amount:"),
    ("Client Rent Amount", "Client Amount:"),
    ("ubh amount", "Total Check Request Amount:"),
    ("Bedroom Sync", "# of Beds"),
    ("Household Size Sync", "Household Size"),
    ("over fmr?", "Over FMR?"),
    ("Has the client been Stepped down?", "Is it Stepdown?"),
]

# Normalized CSV headers for combined "Month/ Year:" (see get_row_values)
PAYMENT_MONTH_CALC_HEADER = "payment month - calc"
PAYMENT_YEAR_CALC_HEADER = "payment year - calc"

# Type of Rental Assistance column (HubSpot export typo "Assitance" + correct spelling)
TYPE_OF_RENTAL_ASSISTANCE_KEYS = (
    "type of rental assitance",
    "type of rental assistance",
)

# Fields that should only be replaced once (e.g. Date: appears in header and signature)
REPLACE_FIRST_ONLY = {"Date:"}

# Fields where label is in one cell, value goes in the adjacent (next) cell
VALUE_IN_NEXT_CELL = {"Total Monthly Rent:", "UBH Rent Amount:", "Client Amount:"}

# Word labels whose CSV values are shown with a leading $
CURRENCY_TEMPLATE_LABELS = frozenset(
    {
        "Total Monthly Rent:",
        "UBH Rent Amount:",
        "Client Amount:",
        "Total Check Request Amount:",
    }
)

# Word labels: value (or blank) is always underlined after the label
UNDERLINE_VALUE_TEMPLATE_LABELS = frozenset(
    {"# of Beds", "Household Size", "Over FMR?", "Is it Stepdown?"}
)

# Visible blank for underlined empty fields (Word shows underline on non-breaking spaces)
UNDERLINE_BLANK_PLACEHOLDER = "\u00a0\u00a0\u00a0"


def normalize_column_name(name: str) -> str:
    """Normalize CSV column header for matching."""
    return name.strip().lower()


def build_column_map(csv_headers: list[str]) -> dict[str, str]:
    """Map normalized column names to original column names."""
    return {normalize_column_name(h): h for h in csv_headers}


def format_payment_date_mm_dd_yyyy(raw: str) -> str:
    """Normalize payment date strings to mm-dd-yyyy for the Word Date: field."""
    s = (raw or "").strip()
    if not s:
        return ""
    if " " in s:
        s = s.split()[0]
    for fmt in ("%Y-%m-%d", "%m/%d/%Y", "%m/%d/%y", "%m-%d-%Y", "%Y/%m/%d"):
        try:
            d = datetime.strptime(s, fmt).date()
            return d.strftime("%m-%d-%Y")
        except ValueError:
            continue
    try:
        d = date.fromisoformat(s)
        return d.strftime("%m-%d-%Y")
    except ValueError:
        return raw


def format_currency_display(raw: str) -> str:
    """Prefix dollar amounts with $ when non-empty; avoid double $."""
    s = (raw or "").strip()
    if not s:
        return ""
    if s.startswith("$"):
        return s
    return f"${s}"


def _underline_display_text(value: str) -> str:
    """Text to show in an underlined run (blank placeholder if empty)."""
    s = (value or "").strip()
    return s if s else UNDERLINE_BLANK_PLACEHOLDER


def _format_month_year_line(month: str, year: str) -> str:
    """Format Payment Month + Payment Year for the Word label 'Month/ Year:'."""
    m = (month or "").strip()
    y = (year or "").strip()
    if m and y:
        return f"{m}, {y}"
    if m:
        return m
    if y:
        return y
    return ""


def get_row_values(row: dict, column_map: dict[str, str]) -> dict[str, str]:
    """Extract values for each field from a row, using flexible column matching."""
    result = {}
    for csv_header, template_label in FIELD_MAPPING:
        # column_map keys are normalized; allow Title Case etc. in FIELD_MAPPING
        orig_col = column_map.get(normalize_column_name(csv_header))
        if orig_col is not None:
            val = row.get(orig_col, "")
            result[template_label] = (val or "").strip()

    month_col = column_map.get(PAYMENT_MONTH_CALC_HEADER)
    year_col = column_map.get(PAYMENT_YEAR_CALC_HEADER)
    if month_col is not None or year_col is not None:
        m_val = (row.get(month_col, "") or "").strip() if month_col else ""
        y_val = (row.get(year_col, "") or "").strip() if year_col else ""
        result["Month/ Year:"] = _format_month_year_line(m_val, y_val)

    for label in CURRENCY_TEMPLATE_LABELS:
        if label in result:
            result[label] = format_currency_display(result[label])

    if result.get("Date:"):
        result["Date:"] = format_payment_date_mm_dd_yyyy(result["Date:"])

    return result


def _norm_compare_str(s: str) -> str:
    """Lowercase stripped string for branch comparisons."""
    return (s or "").strip().casefold()


def _rental_assistance_csv_column(column_map: dict[str, str]) -> str | None:
    """Original CSV header for Type of Rental Assistance, if present."""
    for key in TYPE_OF_RENTAL_ASSISTANCE_KEYS:
        orig = column_map.get(key)
        if orig:
            return orig
    return None


def _get_check_type_and_assistance(
    row: dict, column_map: dict[str, str]
) -> tuple[str, str]:
    """
    Return (check_type_norm, rental_assistance_norm) for branching.
    Missing columns yield empty strings.
    """
    ct_col = column_map.get("check type")
    ra_col = _rental_assistance_csv_column(column_map)
    ct_raw = row.get(ct_col, "") if ct_col else ""
    ra_raw = row.get(ra_col, "") if ra_col else ""
    return _norm_compare_str(ct_raw), _norm_compare_str(ra_raw)


def _is_admin_fee_pest_control(assistance_norm: str) -> bool:
    """Match 'Admin Fee+Pest Control' with flexible spaces around +."""
    s = (assistance_norm or "").strip().casefold()
    s = re.sub(r"\s*\+\s*", "+", s)
    s = re.sub(r"\s+", "", s)
    return s == "adminfee+pestcontrol"


def _is_attorney_fees(assistance_norm: str) -> bool:
    """Match Attorney Fees (and common variants)."""
    s = (assistance_norm or "").strip().casefold()
    return s in ("attorney fees", "attorney fee") or s.startswith("attorney fees")


def _formatted_ubh_from_csv(row: dict, column_map: dict[str, str]) -> str:
    """UBH Amount from CSV with $ formatting (independent of values dict)."""
    col = column_map.get("ubh amount")
    if not col:
        return ""
    return format_currency_display((row.get(col) or "").strip())


def _should_suppress_main_rent_amount_fields(assistance_norm: str) -> bool:
    """
    Do not fill Total Monthly Rent, UBH Rent Amount, Client Amount (and total
    check request) for these Type of Rental Assistance values.
    """
    a = (assistance_norm or "").strip().casefold()
    if a in ("late fee", "utility deposit", "utilities"):
        return True
    if _is_attorney_fees(a):
        return True
    if _is_admin_fee_pest_control(a):
        return True
    return False


def collect_all_tables(doc: Document) -> list:
    """All tables in document order, including nested tables in cells."""

    out: list = []

    def walk_tables(tables) -> None:
        for tbl in tables:
            out.append(tbl)
            for r in tbl.rows:
                for c in r.cells:
                    walk_tables(c.tables)

    walk_tables(doc.tables)
    return out


def find_type_of_assistance_table(doc: Document):
    """
    Table that contains Utility Deposit (and Type of Assistance when possible).

    Prefers a table that has both phrases; otherwise the first table with a
    Utility Deposit row (including nested tables).
    """
    tables = collect_all_tables(doc)
    fallback = None
    for table in tables:
        if find_utility_deposit_row_index(table) is None:
            continue
        if fallback is None:
            fallback = table
        for row in table.rows:
            joined = " ".join(c.text for c in row.cells).lower()
            if "type of assistance" in joined:
                return table
    return fallback


def find_utility_deposit_row_index(table) -> int | None:
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            if "utility deposit" in cell.text.lower():
                return i
    return None


def insert_row_after(table, row_idx: int) -> int:
    """
    Deep-copy the row at row_idx, insert the copy immediately after it, clear
    cell text in the new row. Returns the index of the new row.

    Column B (Excel-style) is cell index 1 when the table has at least two cells.
    """
    tr = table.rows[row_idx]._tr
    new_tr = deepcopy(tr)
    tr.addnext(new_tr)
    tbl = table._tbl
    trs = tbl.findall(qn("w:tr"))
    new_idx = trs.index(new_tr)
    new_row = table.rows[new_idx]
    for cell in new_row.cells:
        set_cell_text(cell, "")
    return new_idx


def _set_row_column_b_text(table, row_idx: int, text: str) -> None:
    """
    Write text to column B (second logical cell, index 1).

    Matches Excel-style B6 when the inserted row sits under Utility Deposit.
    If the template uses merged cells, indices may need adjustment after visual check.
    """
    row = table.rows[row_idx]
    if len(row.cells) >= 2:
        set_cell_text(row.cells[1], text)
    else:
        set_cell_text(row.cells[0], text)


def apply_arrears_and_assistance_table_rules(
    doc: Document,
    row: dict,
    column_map: dict[str, str],
    values: dict[str, str],
) -> None:
    """
    Check Type + Type of Rental Assistance branches.

    For several assistance types, clears main rent / UBH / client amount fields
    and writes UBH into the Type of Assistance table (Utility Deposit row B5 or
    new rows for Late Fee / Admin / Attorney).

    Rule: assistance type Arrears alone leaves Check Type from CSV.
    """
    check_cf, assist_cf = _get_check_type_and_assistance(row, column_map)
    check_is_arrears = check_cf == "arrears"
    ubh = _formatted_ubh_from_csv(row, column_map)

    if _should_suppress_main_rent_amount_fields(assist_cf):
        values["Total Monthly Rent:"] = ""
        values["UBH Rent Amount:"] = ""
        values["Client Amount:"] = ""
        values["Total Check Request Amount:"] = ""

    table = find_type_of_assistance_table(doc)
    uidx = find_utility_deposit_row_index(table) if table else None

    # B5: Utility Deposit row, column B — full label line for these assistance types
    if assist_cf in ("utility deposit", "utilities"):
        if table is None or uidx is None:
            warnings.warn(
                "Utility Deposit / Utilities: could not find Utility Deposit row; skipping B5.",
                UserWarning,
                stacklevel=2,
            )
        else:
            line = f"Utility Deposit: {ubh}" if ubh else "Utility Deposit:"
            _set_row_column_b_text(table, uidx, line)
        if check_is_arrears and assist_cf == "utilities":
            values["Check Type:"] = "Arrears"

    if assist_cf == "late fee":
        if check_is_arrears:
            values["Check Type:"] = "Arrears"
        if table is None or uidx is None:
            warnings.warn(
                "Late Fee: could not find Utility Deposit row; skipping Late Fee row.",
                UserWarning,
                stacklevel=2,
            )
        else:
            new_idx = insert_row_after(table, uidx)
            line = f"Late Fee: {ubh}" if ubh else "Late Fee:"
            _set_row_column_b_text(table, new_idx, line)

    if _is_attorney_fees(assist_cf):
        if table is None or uidx is None:
            warnings.warn(
                "Attorney Fees: could not find Utility Deposit row; skipping B6 row.",
                UserWarning,
                stacklevel=2,
            )
        else:
            new_idx = insert_row_after(table, uidx)
            line = f"Attorney Fees: {ubh}" if ubh else "Attorney Fees:"
            _set_row_column_b_text(table, new_idx, line)

    if _is_admin_fee_pest_control(assist_cf):
        if check_is_arrears:
            values["Check Type:"] = "Arrears"
        if table is None or uidx is None:
            warnings.warn(
                "Admin Fee+Pest Control: could not find Utility Deposit row; skipping row.",
                UserWarning,
                stacklevel=2,
            )
        else:
            new_idx = insert_row_after(table, uidx)
            line = f"Admin Fee+Pest Control: {ubh}" if ubh else "Admin Fee+Pest Control:"
            _set_row_column_b_text(table, new_idx, line)

    if assist_cf == "arrears":
        pass


def replace_in_paragraph(paragraph, old_text: str, new_text: str) -> bool:
    """Replace old_text with new_text in paragraph. Returns True if replaced."""
    if old_text not in paragraph.text:
        return False
    full = paragraph.text
    new_full = full.replace(old_text, new_text, 1)
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_full
    else:
        paragraph.add_run(new_full)
    return True


def _remove_all_runs(paragraph) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def _paragraph_base_font_size(paragraph):
    for run in paragraph.runs:
        if run.font.size:
            return run.font.size
    return None


def replace_label_with_underlined_value(paragraph, label: str, value: str) -> bool:
    """Replace label with label + space + underlined value (or underlined blank)."""
    full = paragraph.text
    if label not in full:
        return False
    idx = full.find(label)
    before = full[:idx]
    after = full[idx + len(label) :]
    fs = _paragraph_base_font_size(paragraph)
    _remove_all_runs(paragraph)
    if before:
        r = paragraph.add_run(before)
        if fs:
            r.font.size = fs
    r = paragraph.add_run(label)
    if fs:
        r.font.size = fs
    r = paragraph.add_run(" ")
    if fs:
        r.font.size = fs
    u_text = _underline_display_text(value)
    ur = paragraph.add_run(u_text)
    ur.font.underline = WD_UNDERLINE.SINGLE
    if fs:
        ur.font.size = fs
    if after:
        r2 = paragraph.add_run(after)
        if fs:
            r2.font.size = fs
    return True


def _merge_underline_segments(
    segments: list[tuple[str, bool]],
) -> list[tuple[str, bool]]:
    if not segments:
        return segments
    out: list[tuple[str, bool]] = [segments[0]]
    for text, ul in segments[1:]:
        prev_t, prev_u = out[-1]
        if text == "":
            continue
        if prev_u == ul:
            out[-1] = (prev_t + text, prev_u)
        else:
            out.append((text, ul))
    return [(t, u) for t, u in out if t != ""]


def _apply_underline_pattern(
    segments: list[tuple[str, bool]],
    pattern: str,
    build_replacement,
) -> list[tuple[str, bool]]:
    """Split non-underlined segments where pattern matches; insert underlined piece(s)."""
    new: list[tuple[str, bool]] = []
    for text, ul in segments:
        if ul or not re.search(pattern, text):
            new.append((text, ul))
            continue
        last = 0
        for m in re.finditer(pattern, text):
            new.append((text[last : m.start()], False))
            new.extend(build_replacement(m))
            last = m.end()
        new.append((text[last:], False))
    return _merge_underline_segments(new)


def _rebuild_paragraph_from_segments(paragraph, segments: list[tuple[str, bool]]) -> None:
    fs = _paragraph_base_font_size(paragraph)
    _remove_all_runs(paragraph)
    for text, ul in segments:
        if text == "":
            continue
        run = paragraph.add_run(text)
        if ul:
            run.font.underline = WD_UNDERLINE.SINGLE
        if fs:
            run.font.size = fs


def fill_comments_paragraph_underlined(paragraph, values: dict[str, str]) -> None:
    """Rewrite Comments body with fills underlined where requested (incl. client amount line)."""
    beds = (values.get("# of Beds") or "").strip()
    household_size = (values.get("Household Size") or "").strip()
    over_fmr = (values.get("Over FMR?") or "").strip()
    stepdown = (values.get("Is it Stepdown?") or "").strip()
    client_amount = (values.get("Client Amount:") or "").strip()

    segs: list[tuple[str, bool]] = [(paragraph.text, False)]

    segs = _apply_underline_pattern(
        segs,
        r"(family of)(\s+)(\.)",
        lambda m: [
            (m.group(1) + " ", False),
            (_underline_display_text(household_size), True),
            (m.group(3), False),
        ],
    )

    segs = _apply_underline_pattern(
        segs,
        r"(a\s+)([Bb]edroom)",
        lambda m: [
            ("a ", False),
            (_underline_display_text(beds), True),
            (" ", False),
            (m.group(2), False),
        ],
    )

    segs = _apply_underline_pattern(
        segs,
        r"(The amount is )(\S*)( over FMR)",
        lambda m: [
            (m.group(1), False),
            (_underline_display_text(over_fmr), True),
            (m.group(3), False),
        ],
    )

    # Template phrase "amount is … stepdown" -> value underlined between "amount is " and " stepdown"
    new_segs: list[tuple[str, bool]] = []
    for chunk, ul in segs:
        if ul:
            new_segs.append((chunk, ul))
            continue
        pat = r"amount is\s+stepdown"
        m = re.search(pat, chunk)
        if not m:
            new_segs.append((chunk, False))
            continue
        new_segs.append((chunk[: m.start()], False))
        new_segs.append(("amount is ", False))
        new_segs.append((_underline_display_text(stepdown), True))
        new_segs.append((" stepdown", False))
        new_segs.append((chunk[m.end() :], False))
    segs = _merge_underline_segments(new_segs)

    # "The client is responsible for … amount" — Client Rent Amount (underlined)
    new_segs = []
    for chunk, ul in segs:
        if ul:
            new_segs.append((chunk, ul))
            continue
        pat = r"(responsible for)(\s+)(amount)"
        m = re.search(pat, chunk)
        if not m:
            new_segs.append((chunk, False))
            continue
        new_segs.append((chunk[: m.start()], False))
        new_segs.append(("responsible for ", False))
        new_segs.append((_underline_display_text(client_amount), True))
        new_segs.append((" amount", False))
        new_segs.append((chunk[m.end() :], False))
    segs = _merge_underline_segments(new_segs)

    _rebuild_paragraph_from_segments(paragraph, segs)


def set_cell_text(cell, text: str, font_size: int = 9) -> None:
    """Set the first paragraph of a cell to the given text with specified font size (default 9pt)."""
    if cell.paragraphs:
        para = cell.paragraphs[0]
        # Clear existing runs and add a fresh run with correct font size
        for run in para.runs:
            run.text = ""
        run = para.add_run(text)
        run.font.size = Pt(font_size)
    else:
        para = cell.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(font_size)


def cell_contains_label(cell, template_label: str) -> bool:
    """Check if cell contains the label (with or without trailing colon)."""
    cell_text = " ".join(p.text for p in cell.paragraphs).strip()
    label_with_colon = template_label
    label_without_colon = template_label.rstrip(":")
    return label_with_colon in cell_text or label_without_colon in cell_text


def process_cell_paragraphs(cell, row, cell_idx: int, values: dict[str, str], replaced_count: dict[str, int]) -> None:
    """Process paragraphs in a cell for replacements."""
    # Check for "value in next cell" fields first
    for template_label in VALUE_IN_NEXT_CELL:
        if template_label not in values:
            continue
        if not cell_contains_label(cell, template_label):
            continue
        # Put value in the next cell (same row, next column)
        if cell_idx + 1 < len(row.cells):
            set_cell_text(row.cells[cell_idx + 1], values[template_label])
        return

    # Standard replacement: "Client Name:" -> "Client Name: John Test" in same cell
    for paragraph in cell.paragraphs:
        for template_label, value in values.items():
            if template_label in VALUE_IN_NEXT_CELL:
                continue

            if template_label in REPLACE_FIRST_ONLY:
                if replaced_count[template_label] > 0:
                    continue

            if template_label in UNDERLINE_VALUE_TEMPLATE_LABELS:
                if replace_label_with_underlined_value(paragraph, template_label, value):
                    if template_label in REPLACE_FIRST_ONLY:
                        replaced_count[template_label] += 1
                    break
                continue

            # Build replacement: "Client Name:" -> "Client Name: John Test"
            replacement = f"{template_label} {value}" if value else template_label

            if replace_in_paragraph(paragraph, template_label, replacement):
                if template_label in REPLACE_FIRST_ONLY:
                    replaced_count[template_label] += 1
                break  # Only one replacement per paragraph


def process_table(table, values: dict[str, str], replaced_count: dict[str, int]) -> None:
    """Process a table and any nested tables recursively."""
    for row in table.rows:
        for cell_idx, cell in enumerate(row.cells):
            process_cell_paragraphs(cell, row, cell_idx, values, replaced_count)
            for nested_table in cell.tables:
                process_table(nested_table, values, replaced_count)


def insert_signature_tag_in_director_cell(doc: Document) -> bool:
    """
    Append the PandaDoc signature field tag (white text) to the Director
    signature cell. Returns True if the cell was found and tagged.
    """
    for table in collect_all_tables(doc):
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip().startswith("Director:"):
                    para = cell.paragraphs[0]
                    run = para.add_run(f" {PANDADOC_SIGNATURE_TAG}")
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    return True
    return False


def fill_template(
    doc: Document,
    values: dict[str, str],
    row: dict,
    column_map: dict[str, str],
    include_signature_tag: bool = False,
) -> None:
    """Fill the Word document template with values from a CSV row."""
    apply_arrears_and_assistance_table_rules(doc, row, column_map, values)

    replaced_count = {label: 0 for label in REPLACE_FIRST_ONLY}

    for table in doc.tables:
        process_table(table, values, replaced_count)

    # Comments paragraph: keys are Word labels (same strings as FIELD_MAPPING values / Month/ Year:).
    for para in doc.paragraphs:
        text = para.text
        if "Comments:" in text and "Bedroom" in text and "FMR" in text:
            fill_comments_paragraph_underlined(para, values)
            break

    if include_signature_tag:
        if not insert_signature_tag_in_director_cell(doc):
            warnings.warn(
                "Signature tag requested but no 'Director:' cell found in template.",
                UserWarning,
                stacklevel=2,
            )


def generate_combined_docx(
    rows: list[dict],
    template_path: Path | str,
    output_path: Path | str,
    include_signature_tag: bool = False,
) -> Path | None:
    """
    Fill the template once per row dict and merge all forms into one combined
    DOCX at output_path. Rows are plain dicts keyed by the CSV-style headers
    FIELD_MAPPING expects (e.g. the output of hubspot_pull.get_rows_for_batch).

    Returns the output path, or None if rows is empty (no file written).
    """
    template_path = Path(template_path)
    output_path = Path(output_path)

    if not rows:
        return None

    column_map = build_column_map(list(rows[0].keys()))

    docs_to_merge = []
    for row in rows:
        values = get_row_values(row, column_map)
        doc = Document(str(template_path))
        fill_template(
            doc, values, row, column_map, include_signature_tag=include_signature_tag
        )
        docs_to_merge.append(doc)

    master = docs_to_merge[0]
    composer = Composer(master)
    for doc in docs_to_merge[1:]:
        composer.append(doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    composer.save(str(output_path))
    return output_path


def sanitize_filename(name: str) -> str:
    """Remove or replace characters unsafe for filenames."""
    return re.sub(r'[<>:"/\\|?*]', "_", name).strip() or "Unknown"


def _detect_csv_encoding(csv_path: Path) -> str:
    """Use UTF-8 when valid; otherwise common Windows/mac exports (cp1252, latin-1)."""
    data = csv_path.read_bytes()
    for enc in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            data.decode(enc)
            return enc
        except UnicodeDecodeError:
            continue
    return "latin-1"


def _soffice_executable():
    """Path to LibreOffice soffice, if available."""
    exe = shutil.which("soffice") or shutil.which("libreoffice")
    if exe:
        return exe
    if sys.platform == "darwin":
        mac_lo = Path("/Applications/LibreOffice.app/Contents/MacOS/soffice")
        if mac_lo.is_file():
            return str(mac_lo)
    return None


def convert_docx_to_pdf(docx_path: Path, pdf_path: Path) -> bool:
    """
    Create pdf_path from docx_path.

    Tries docx2pdf (uses Microsoft Word on macOS/Windows), then LibreOffice headless.
    """
    docx_path = docx_path.resolve()
    pdf_path = pdf_path.resolve()
    pdf_path.parent.mkdir(parents=True, exist_ok=True)

    try:
        from docx2pdf import convert as docx2pdf_convert

        docx2pdf_convert(str(docx_path), str(pdf_path))
        if pdf_path.is_file() and pdf_path.stat().st_size > 0:
            return True
    except Exception:
        pass

    soffice = _soffice_executable()
    if not soffice:
        return False

    outdir = pdf_path.parent
    try:
        subprocess.run(
            [
                soffice,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(outdir),
                str(docx_path),
            ],
            check=True,
            capture_output=True,
            timeout=300,
        )
    except (subprocess.CalledProcessError, subprocess.TimeoutExpired, OSError):
        return False

    produced = outdir / f"{docx_path.stem}.pdf"
    if not produced.is_file() or produced.stat().st_size == 0:
        return False
    if produced.resolve() != pdf_path.resolve():
        shutil.move(str(produced), str(pdf_path))
    return pdf_path.is_file()


def main():
    parser = argparse.ArgumentParser(
        description="Generate Word check request forms from CSV data."
    )
    parser.add_argument(
        "--data-dir",
        default="Data",
        help="Folder containing the single CSV file (default: Data)",
    )
    parser.add_argument(
        "--template",
        default="Form Template/Rapid Rehousing Program Check Request Form - Template.docx",
        help="Path to Word template",
    )
    parser.add_argument(
        "--output",
        default="Output",
        help="Output directory for generated forms",
    )
    parser.add_argument(
        "--no-pdf",
        action="store_true",
        help="Skip PDF export (only write the combined .docx)",
    )
    args = parser.parse_args()

    script_dir = Path(__file__).parent.resolve()
    data_dir = script_dir / args.data_dir
    template_path = script_dir / args.template
    output_dir = script_dir / args.output

    if not data_dir.is_dir():
        print(f"Error: Data folder not found: {data_dir}")
        return 1
    csv_files = list(data_dir.glob("*.csv"))
    if not csv_files:
        print("Error: No CSV file found in Data folder.")
        return 1
    if len(csv_files) > 1:
        print("Error: More than one CSV file in Data folder. Keep only one.")
        return 1
    csv_path = csv_files[0]

    if not template_path.exists():
        print(f"Error: Template not found: {template_path}")
        return 1

    output_dir.mkdir(parents=True, exist_ok=True)

    with open(csv_path, newline="", encoding=_detect_csv_encoding(csv_path)) as f:
        rows = list(csv.DictReader(f))

    out_path = output_dir / "Check_Requests_Combined.docx"
    if generate_combined_docx(rows, template_path, out_path) is None:
        print("No rows in CSV. Nothing to generate.")
        return 0

    print(f"\nGenerated Word: {out_path} ({len(rows)} form(s))")

    if args.no_pdf:
        return 0

    pdf_path = output_dir / "Check_Requests_Combined.pdf"
    if convert_docx_to_pdf(out_path, pdf_path):
        print(f"Generated PDF:  {pdf_path}")
    else:
        print(
            "\nWarning: Could not create PDF. Install Microsoft Word and `pip install docx2pdf`, "
            "or install LibreOffice (soffice on PATH, or the app under /Applications on macOS).",
            file=sys.stderr,
        )

    return 0


if __name__ == "__main__":
    exit(main())
